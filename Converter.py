import os
import sys
import pathlib
import av
import make_key
import traceback
from PIL import Image
import rawpy
from pdf2docx import Converter
from docx import Document
from odf.opendocument import load as load_odt
from odf.opendocument import OpenDocumentText
from odf.text import P
from odf import teletype
from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
import pandas as pd
import xlwt

file_types = {
    # video files
    ".mp4": [
        ("convert to mp3", "Convert To MP3", "MP3"),
        ("convert to wav", "Convert To WAV", "WAV"),
        ("convert to flac", "Convert To FLAC", "FLAC"),
        ("convert to audio ogg", "Convert To Audio OGG", "OGG"),
        ("convert to opus", "Convert To OPUS", "OPUS"),
        ("convert to mkv", "Convert To MKV", "MKV"),
        ("convert to mov", "Convert To MOV", "MOV"),
        ("convert to avi", "Convert To AVI", "AVI"),
        ("convert to webm", "Convert To WEBM", "WEBM")
    ],
    ".mkv": [
        ("convert to mp3", "Convert To MP3", "MP3"),
        ("convert to flac", "Convert To FLAC", "FLAC"),
        ("convert to audio ogg", "Convert To Audio OGG", "OGG"),
        ("convert to opus", "Convert To OPUS", "OPUS"),
        ("convert to wav", "Convert To WAV", "WAV"),
        ("convert to mp4", "Convert To MP4", "MP4"),
        ("convert to mov", "Convert To MOV", "MOV"),
        ("convert to avi", "Convert To AVI", "AVI"),
        ("convert to webm", "Convert To WEBM", "WEBM")
    ],
    ".mov": [
        ("convert to mp3", "Convert To MP3", "MP3"),
        ("convert to flac", "Convert To FLAC", "FLAC"),
        ("convert to audio ogg", "Convert To Audio OGG", "OGG"),
        ("convert to opus", "Convert To OPUS", "OPUS"),
        ("convert to wav", "Convert To WAV", "WAV"),
        ("convert to mp4", "Convert To MP4", "MP4"),
        ("convert to avi", "Convert To AVI", "AVI"),
        ("convert to webm", "Convert To WEBM", "WEBM"),
        ("convert to mkv", "Convert To MKV", "MKV")
    ],
    ".avi": [
        ("convert to mp3", "Convert To MP3", "MP3"),
        ("convert to flac", "Convert To FLAC", "FLAC"),
        ("convert to audio ogg", "Convert To Audio OGG", "OGG"),
        ("convert to opus", "Convert To OPUS", "OPUS"),
        ("convert to mp4", "Convert To MP4", "MP4"),
        ("convert to wav", "Convert To WAV", "WAV"),
        ("convert to webm", "Convert To WEBM", "WEBM"),
        ("convert to mkv", "Convert To MKV", "MKV"),
        ("convert to mov", "Convert To MOV", "MOV")
    ],
    ".webm": [
        ("convert to mp3", "Convert To MP3", "MP3"),
        ("convert to flac", "Convert To FLAC", "FLAC"),
        ("convert to audio ogg", "Convert To Audio OGG", "OGG"),
        ("convert to opus", "Convert To OPUS", "OPUS"),
        ("convert to wav", "Convert To WAV", "WAV"),
        ("convert to mp4", "Convert To MP4", "MP4"),
        ("convert to avi", "Convert To AVI", "AVI"),
        ("convert to mkv", "Convert To MKV", "MKV"),
        ("convert to mov", "Convert To MOV", "MOV")
    ],
    # audio files
    ".mp3": [
        ("convert to mp3", "Convert To MP3", "MP3"),
        ("convert to flac", "Convert To FLAC", "FLAC"),
        ("convert to audio ogg", "Convert To Audio OGG", "OGG"),
        ("convert to wav", "Convert To WAV", "WAV"),
        ("convert to opus", "Convert To OPUS", "OPUS")
    ],
    ".wav": [
        ("convert to mp3", "Convert To MP3", "MP3"),
        ("convert to flac", "Convert To FLAC", "FLAC"),
        ("convert to audio ogg", "Convert To Audio OGG", "OGG"),
        ("convert to opus", "Convert To OPUS", "OPUS")
    ],
    ".ogg": [
        ("convert to mp3", "Convert To MP3", "MP3"),
        ("convert to flac", "Convert To FLAC", "FLAC"),
        ("convert to wav", "Convert To WAV", "WAV"),
        ("convert to opus", "Convert To OPUS", "OPUS")
    ],
    ".flac": [
        ("convert to mp3", "Convert To MP3", "MP3"),
        ("convert to wav", "Convert To WAV", "WAV"),
        ("convert to audio ogg", "Convert To Audio OGG", "OGG"),
        ("convert to opus", "Convert To OPUS", "OPUS")
    ],
    ".opus": [
        ("convert to mp3", "Convert To MP3", "MP3"),
        ("convert to wav", "Convert To WAV", "WAV"),
        ("convert to audio ogg", "Convert To Audio OGG", "OGG"),
        ("convert to flac", "Convert To FLAC", "FLAC"),
    ],
    # images
    ".png": [
        ("convert to jpg", "Convert To JPG", "JPG"),
        ("convert to jpeg", "Convert To JPEG", "JPEG"),
        ("convert to webp", "Convert To WEBP", "WEBP"),
        ("convert to ico", "Convert To ICO", "ICO"),
        ("convert to pdf", "Convert To PDF", "PDF")
    ],
    ".jpg": [
        ("convert to png", "Convert To PNG", "PNG"),
        ("convert to jpeg", "Convert To JPEG", "JPEG"),
        ("convert to webp", "Convert To WEBP", "WEBP"),
        ("convert to ico", "Convert To ICO", "ICO"),
        ("convert to pdf", "Convert To PDF", "PDF")
    ],
    ".jpeg": [
        ("convert to png", "Convert To PNG", "PNG"),
        ("convert to jpg", "Convert To JPG", "JPG"),
        ("convert to webp", "Convert To WEBP", "WEBP"),
        ("convert to ico", "Convert To ICO", "ICO"),
        ("convert to pdf", "Convert To PDF", "PDF")
    ],
    ".webp": [
        ("convert to png", "Convert To PNG", "PNG"),
        ("convert to jpg", "Convert To JPG", "JPG"),
        ("convert to jpeg", "Convert To JPEG", "JPEG"),
        ("convert to ico", "Convert To ICO", "ICO"),
        ("convert to pdf", "Convert To PDF", "PDF")
    ],
    ".ico": [
        ("convert to png", "Convert To PNG", "PNG"),
        ("convert to jpg", "Convert To JPG", "JPG"),
        ("convert to jpeg", "Convert To JPEG", "JPEG"),
        ("convert to webp", "Convert To WEBP", "WEBP"),
        ("convert to pdf", "Convert To PDF", "PDF")
    ],
    ".raw": [
        ("convert to png", "Convert To PNG", "PNG"),
        ("convert to jpg", "Convert To JPG", "JPG"),
        ("convert to jpeg", "Convert To JPEG", "JPEG"),
        ("convert to webp", "Convert To WEBP", "WEBP"),
        ("convert to ico", "Convert To ICO", "ICO"),
        ("convert to pdf", "Convert To PDF", "PDF")
    ],
    # documents
    ".pdf": [
        ("convert to docx", "Convert To DOCX", "DOCXFPDF")
    ],
    ".docx": [
        ("convert to pdf", "Convert To PDF", "DOCPDF"),
        ("convert to txt", "Convert To TXT", "DOCTXT"),
        ("convert to odt", "Convert To ODT", "DOCODT")
    ],
    ".txt": [
        ("convert to pdf", "Convert To PDF", "DOCPDF"),
        ("convert to odt", "Convert To ODT", "DOCODT"),
        ("convert to docx", "Convert To DOCX", "DOCDOCX")
    ],
    ".odt": [
        ("convert to pdf", "Convert To PDF", "DOCPDF"),
        ("convert to txt", "Convert To TXT", "DOCTXT"),
        ("convert to docx", "Convert To DOCX", "DOCDOCX")
    ],

    # spreadsheets
    ".xlsx": [
        ("convert to pdf", "Convert To PDF", "SHEETPDF"),
        ("convert to xls", "Convert To XLS", "SHEETXLS"),
        ("convert to ods", "Convert To ODS", "SHEETODS"),
        ("convert to csv", "Convert To CSV", "SHEETCSV"),
        ("convert to tsv", "Convert To TSV", "SHEETTSV")
    ],
    ".xls": [
        ("convert to pdf", "Convert To PDF", "SHEETPDF"),
        ("convert to xlsx", "Convert To XLSX", "SHEETXLSX"),
        ("convert to ods", "Convert To ODS", "SHEETODS"),
        ("convert to csv", "Convert To CSV", "SHEETCSV"),
        ("convert to tsv", "Convert To TSV", "SHEETTSV")
    ],
    ".ods": [
        ("convert to pdf", "Convert To PDF", "SHEETPDF"),
        ("convert to xlsx", "Convert To XLSX", "SHEETXLSX"),
        ("convert to xls", "Convert To XLS", "SHEETXLS"),
        ("convert to csv", "Convert To CSV", "SHEETCSV"),
        ("convert to tsv", "Convert To TSV", "SHEETTSV")
    ],
    ".csv": [
        ("convert to pdf", "Convert To PDF", "SHEETPDF"),
        ("convert to xlsx", "Convert To XLSX", "SHEETXLSX"),
        ("convert to xls", "Convert To XLS", "SHEETXLS"),
        ("convert to ods", "Convert To ODS", "SHEETODS"),
        ("convert to tsv", "Convert To TSV", "SHEETTSV")
    ],
    ".xlsb": [
        ("convert to pdf", "Convert To PDF", "SHEETPDF"),
        ("convert to xlsx", "Convert To XLSX", "SHEETXLSX"),
        ("convert to xls", "Convert To XLS", "SHEETXLS"),
        ("convert to ods", "Convert To ODS", "SHEETODS"),
        ("convert to csv", "Convert To CSV", "SHEETCSV"),
        ("convert to tsv", "Convert To TSV", "SHEETTSV")
    ],
    ".xlsm": [
        ("convert to pdf", "Convert To PDF", "SHEETPDF"),
        ("convert to xlsx", "Convert To XLSX", "SHEETXLSX"),
        ("convert to xls", "Convert To XLS", "SHEETXLS"),
        ("convert to ods", "Convert To ODS", "SHEETODS"),
        ("convert to csv", "Convert To CSV", "SHEETCSV"),
        ("convert to tsv", "Convert To TSV", "SHEETTSV")
    ],
    ".tsv": [
        ("convert to pdf", "Convert To PDF", "SHEETPDF"),
        ("convert to xlsx", "Convert To XLSX", "SHEETXLSX"),
        ("convert to xls", "Convert To XLS", "SHEETXLS"),
        ("convert to ods", "Convert To ODS", "SHEETODS"),
        ("convert to csv", "Convert To CSV", "SHEETCSV")
    ]
}

av.logging.set_level(av.logging.VERBOSE)

def ConvertFile(file_path, convert_type):
    input_file = None
    output_file = None
    
    output_file_pre_suffix = file_path.removesuffix(pathlib.Path(file_path).suffix)
    output_file_path = output_file_pre_suffix + '.' + convert_type.lower()
    
    try:     
        match convert_type.lower():
            
            # video file conversions
            
            case "mp4" | "mkv" | "mov" :
                input_file = av.open(file_path)
                output_file = av.open(output_file_path, 'w')
                Remux(input_file, output_file)
                pass
            case "webm":
                input_file = av.open(file_path)
                output_file = av.open(output_file_path, 'w')
                Transcode(input_file, output_file, "vp8", "libopus")
                pass
            case "avi":
                input_file = av.open(file_path)
                output_file = av.open(output_file_path, 'w')
                Transcode(input_file, output_file, "libsvtav1", "libmp3lame")
                pass
            
            # audio file conversions
            
            case "mp3" | "wav" | "flac" | "ogg" | "opus":
                input_file = av.open(file_path)
                output_file = av.open(output_file_path, "w")

                audio_codecs = {
                    "mp3": "libmp3lame",
                    "wav": "pcm_s16le",
                    "flac": "flac",
                    "ogg": "vorbis",
                    "opus": "libopus"
                }

                TranscodeAudio(
                    input_file,
                    output_file,
                    audio_codecs[convert_type.lower()]
                )
            
            # image file conversions
            
            case "png" | "jpg" | "jpeg" | "webp" | "ico" | "pdf":
                if pathlib.Path(file_path).suffix.lower() == ".raw":
                    with rawpy.imread(file_path) as raw_image:
                        rgb = raw_image.postprocess()
                        image = Image.fromarray(rgb)
                else:
                    image = Image.open(file_path)

                try:
                    if convert_type.lower() in ("jpg", "jpeg"):
                        if image.mode in ("RGBA", "LA"):
                            background = Image.new(
                                "RGB",
                                image.size,
                                (255, 255, 255)
                            )

                            alpha = image.getchannel("A")

                            background.paste(
                                image,
                                mask=alpha
                            )

                            image = background

                        elif image.mode != "RGB":
                            image = image.convert("RGB")

                    elif convert_type.lower() == "pdf":
                        if image.mode in ("RGBA", "LA", "P"):
                            background = Image.new(
                                "RGB",
                                image.size,
                                (255, 255, 255)
                            )

                            if image.mode == "P":
                                image = image.convert("RGBA")

                            if image.mode in ("RGBA", "LA"):
                                background.paste(
                                    image,
                                    mask=image.getchannel("A")
                                )
                                image = background
                            else:
                                image = image.convert("RGB")

                    image.save(output_file_path)
                finally:
                    image.close()
            
            # document conversions

            case "docxfpdf":
                ConvertPdfToDocx(
                    file_path,
                    output_file_pre_suffix + ".docx"
                )

            case "docpdf":
                ConvertDocument(
                    file_path,
                    output_file_pre_suffix + ".pdf",
                    "pdf"
                )

            case "doctxt":
                ConvertDocument(
                    file_path,
                    output_file_pre_suffix + ".txt",
                    "txt"
                )

            case "docodt":
                ConvertDocument(
                    file_path,
                    output_file_pre_suffix + ".odt",
                    "odt"
                )

            case "docdocx":
                ConvertDocument(
                    file_path,
                    output_file_pre_suffix + ".docx",
                    "docx"
                )

            # spreadsheet conversions

            case "sheetpdf":
                ConvertSpreadsheet(
                    file_path,
                    output_file_pre_suffix + ".pdf",
                    "pdf"
                )

            case "sheetxlsx":
                ConvertSpreadsheet(
                    file_path,
                    output_file_pre_suffix + ".xlsx",
                    "xlsx"
                )

            case "sheetxls":
                ConvertSpreadsheet(
                    file_path,
                    output_file_pre_suffix + ".xls",
                    "xls"
                )

            case "sheetods":
                ConvertSpreadsheet(
                    file_path,
                    output_file_pre_suffix + ".ods",
                    "ods"
                )

            case "sheetcsv":
                ConvertSpreadsheet(
                    file_path,
                    output_file_pre_suffix + ".csv",
                    "csv"
                )

            case "sheettsv":
                ConvertSpreadsheet(
                    file_path,
                    output_file_pre_suffix + ".tsv",
                    "tsv"
                )

            case _:
                pass
    finally:
        if input_file is not None:
            input_file.close()

        if output_file is not None:
            output_file.close()

def ReadDocumentText(file_path):
    extension = pathlib.Path(file_path).suffix.lower()

    if extension == ".txt":
        return pathlib.Path(file_path).read_text(
            encoding="utf-8",
            errors="replace"
        )

    if extension == ".docx":
        document = Document(file_path)
        return "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
        )

    if extension == ".odt":
        document = load_odt(file_path)
        paragraphs = document.getElementsByType(P)

        return "\n".join(
            teletype.extractText(paragraph)
            for paragraph in paragraphs
        )

    raise ValueError(
        "Unsupported document input format: " + extension
    )


def SaveTextAsPdf(text, output_file_path):
    pdf = canvas.Canvas(output_file_path, pagesize=A4)
    page_width, page_height = A4

    left_margin = 50
    top_margin = 50
    bottom_margin = 50
    line_height = 14
    max_characters = 95

    y = page_height - top_margin

    for original_line in text.splitlines() or [""]:
        line = original_line

        while len(line) > max_characters:
            split_at = line.rfind(" ", 0, max_characters)

            if split_at <= 0:
                split_at = max_characters

            part = line[:split_at]
            pdf.drawString(left_margin, y, part)
            y -= line_height
            line = line[split_at:].lstrip()

            if y < bottom_margin:
                pdf.showPage()
                y = page_height - top_margin

        pdf.drawString(left_margin, y, line)
        y -= line_height

        if y < bottom_margin:
            pdf.showPage()
            y = page_height - top_margin

    pdf.save()


def ConvertDocument(file_path, output_file_path, output_format):
    text = ReadDocumentText(file_path)

    if output_format == "txt":
        pathlib.Path(output_file_path).write_text(
            text,
            encoding="utf-8"
        )

    elif output_format == "docx":
        document = Document()

        for line in text.splitlines():
            document.add_paragraph(line)

        document.save(output_file_path)

    elif output_format == "odt":
        document = OpenDocumentText()

        for line in text.splitlines():
            paragraph = P()
            paragraph.addText(line)
            document.text.addElement(paragraph)

        document.save(output_file_path)

    elif output_format == "pdf":
        SaveTextAsPdf(text, output_file_path)

    else:
        raise ValueError(
            "Unsupported document output format: " + output_format
        )


def ConvertPdfToDocx(file_path, output_file_path):
    converter = Converter(file_path)

    try:
        converter.convert(output_file_path)
    finally:
        converter.close()


def ReadSpreadsheet(file_path):
    extension = pathlib.Path(file_path).suffix.lower()

    if extension == ".csv":
        return {"Sheet1": pd.read_csv(file_path)}

    if extension == ".tsv":
        return {
            "Sheet1": pd.read_csv(
                file_path,
                sep="\t"
            )
        }

    engines = {
        ".xlsx": "openpyxl",
        ".xlsm": "openpyxl",
        ".xls": "xlrd",
        ".ods": "odf",
        ".xlsb": "pyxlsb"
    }

    engine = engines.get(extension)

    if engine is None:
        raise ValueError(
            "Unsupported spreadsheet input format: " + extension
        )

    return pd.read_excel(
        file_path,
        sheet_name=None,
        engine=engine
    )


def SaveSpreadsheetAsXls(sheets, output_file_path):
    workbook = xlwt.Workbook()

    for sheet_name, dataframe in sheets.items():
        safe_name = str(sheet_name)[:31] or "Sheet"
        worksheet = workbook.add_sheet(safe_name)

        for column_index, column_name in enumerate(dataframe.columns):
            worksheet.write(0, column_index, str(column_name))

        for row_index, row in enumerate(
            dataframe.itertuples(index=False),
            start=1
        ):
            for column_index, value in enumerate(row):
                if pd.isna(value):
                    value = ""

                worksheet.write(
                    row_index,
                    column_index,
                    value
                )

    workbook.save(output_file_path)


def SaveSpreadsheetAsPdf(sheets, output_file_path):
    styles = getSampleStyleSheet()

    document = SimpleDocTemplate(
        output_file_path,
        pagesize=landscape(A4),
        leftMargin=24,
        rightMargin=24,
        topMargin=24,
        bottomMargin=24
    )

    elements = []

    for sheet_name, dataframe in sheets.items():
        elements.append(
            Paragraph(str(sheet_name), styles["Heading2"])
        )

        table_data = [
            [str(column) for column in dataframe.columns]
        ]

        for row in dataframe.itertuples(index=False):
            table_data.append(
                [
                    "" if pd.isna(value) else str(value)
                    for value in row
                ]
            )

        table = Table(
            table_data,
            repeatRows=1
        )

        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("FONTSIZE", (0, 0), (-1, -1), 7)
                ]
            )
        )

        elements.append(table)
        elements.append(Spacer(1, 18))

    document.build(elements)


def ConvertSpreadsheet(file_path, output_file_path, output_format):
    sheets = ReadSpreadsheet(file_path)

    if output_format == "csv":
        first_sheet = next(iter(sheets.values()))
        first_sheet.to_csv(
            output_file_path,
            index=False
        )

    elif output_format == "tsv":
        first_sheet = next(iter(sheets.values()))
        first_sheet.to_csv(
            output_file_path,
            index=False,
            sep="\t"
        )

    elif output_format == "xlsx":
        with pd.ExcelWriter(
            output_file_path,
            engine="openpyxl"
        ) as writer:
            for sheet_name, dataframe in sheets.items():
                dataframe.to_excel(
                    writer,
                    sheet_name=str(sheet_name)[:31],
                    index=False
                )

    elif output_format == "ods":
        with pd.ExcelWriter(
            output_file_path,
            engine="odf"
        ) as writer:
            for sheet_name, dataframe in sheets.items():
                dataframe.to_excel(
                    writer,
                    sheet_name=str(sheet_name)[:31],
                    index=False
                )

    elif output_format == "xls":
        SaveSpreadsheetAsXls(
            sheets,
            output_file_path
        )

    elif output_format == "pdf":
        SaveSpreadsheetAsPdf(
            sheets,
            output_file_path
        )

    else:
        raise ValueError(
            "Unsupported spreadsheet output format: " + output_format
        )


def Remux(input_file, output_file):
    stream_map = {}

    for in_stream in input_file.streams:
        if in_stream.type not in ("video", "audio"):
            continue

        out_stream = output_file.add_stream_from_template(in_stream)
        stream_map[in_stream.index] = out_stream

    for packet in input_file.demux():
        if packet.dts is None:
            continue

        if packet.stream.index not in stream_map:
            continue

        packet.stream = stream_map[packet.stream.index]
        output_file.mux(packet)

def Transcode(input_file, output_file, encoderVideo, encoderAudio):
    stream_map = {}

    for in_stream in input_file.streams:
        if in_stream.type == "video":
            out_stream = output_file.add_stream(
                encoderVideo,
                rate=in_stream.average_rate
            )

            out_stream.width = in_stream.codec_context.width
            out_stream.height = in_stream.codec_context.height
            out_stream.pix_fmt = "yuv420p"

            stream_map[in_stream.index] = out_stream

        elif in_stream.type == "audio":
            out_stream = output_file.add_stream(
                encoderAudio,
                rate=in_stream.codec_context.sample_rate
            )

            out_stream.layout = in_stream.codec_context.layout
            stream_map[in_stream.index] = out_stream

    for packet in input_file.demux():
        if packet.stream.index not in stream_map:
            continue

        out_stream = stream_map[packet.stream.index]

        for frame in packet.decode():
            for output_packet in out_stream.encode(frame):
                output_file.mux(output_packet)

    for out_stream in stream_map.values():
        for packet in out_stream.encode():
            output_file.mux(packet)

def TranscodeAudio(input_file, output_file, encoder):
    input_stream = next(
        (
            stream
            for stream in input_file.streams
            if stream.type == "audio"
        ),
        None
    )

    if input_stream is None:
        raise ValueError("The input file contains no audio stream")

    sample_rate = input_stream.codec_context.sample_rate or 48000

    output_stream = output_file.add_stream(
        encoder,
        rate=sample_rate
    )

    if encoder == "libmp3lame":
        output_stream.bit_rate = 320000

    elif encoder == "vorbis":
        output_stream.codec_context.options = {
            "strict": "-2"
        }
        output_stream.bit_rate = 192000

    elif encoder == "libopus":
        output_stream.bit_rate = 192000

    for frame in input_file.decode(input_stream):
        for packet in output_stream.encode(frame):
            output_file.mux(packet)

    # Flush remaining encoded audio
    for packet in output_stream.encode():
        output_file.mux(packet)

if __name__ == "__main__":
    is_uninstalling = (
        len(sys.argv) > 1
        and sys.argv[1] == "--uninstall"
    )

    try:
        if is_uninstalling:
            make_key.RemoveExtensions(file_types)

        elif len(sys.argv) > 2:
            ConvertFile(sys.argv[1], sys.argv[2])

        else:
            make_key.CreateExtensions(file_types)

    except Exception:
        traceback.print_exc()

        if not is_uninstalling:
            input("\nPress Enter to close...")
