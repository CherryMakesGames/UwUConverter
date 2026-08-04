import pathlib

from docx import Document
from odf.opendocument import load as load_odt
from odf.opendocument import OpenDocumentText
from odf.text import P
from odf import teletype
from pdf2docx import Converter
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def convert_document(
    file_path,
    output_file_path,
    output_format
):
    if output_format == "docx_from_pdf":
        convert_pdf_to_docx(
            file_path,
            output_file_path
        )
        return

    text = read_document_text(file_path)

    if output_format == "txt":
        pathlib.Path(output_file_path).write_text(
            text,
            encoding="utf-8"
        )

    elif output_format == "docx":
        document = Document()

        for line in text.splitlines():
            document.add_paragraph(line)

        document.save(output_file_path)

    elif output_format == "odt":
        document = OpenDocumentText()

        for line in text.splitlines():
            paragraph = P()
            paragraph.addText(line)
            document.text.addElement(paragraph)

        document.save(output_file_path)

    elif output_format == "pdf":
        save_text_as_pdf(
            text,
            output_file_path
        )

    else:
        raise ValueError(
            "Unsupported document output format: "
            + output_format
        )


def read_document_text(file_path):
    extension = pathlib.Path(file_path).suffix.lower()

    if extension == ".txt":
        return pathlib.Path(file_path).read_text(
            encoding="utf-8",
            errors="replace"
        )

    if extension == ".docx":
        document = Document(file_path)

        return "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
        )

    if extension == ".odt":
        document = load_odt(file_path)
        paragraphs = document.getElementsByType(P)

        return "\n".join(
            teletype.extractText(paragraph)
            for paragraph in paragraphs
        )

    raise ValueError(
        "Unsupported document input format: "
        + extension
    )


def save_text_as_pdf(text, output_file_path):
    pdf = canvas.Canvas(
        output_file_path,
        pagesize=A4
    )

    _, page_height = A4

    left_margin = 50
    top_margin = 50
    bottom_margin = 50
    line_height = 14
    max_characters = 95

    y = page_height - top_margin

    for original_line in text.splitlines() or [""]:
        line = original_line

        while len(line) > max_characters:
            split_at = line.rfind(
                " ",
                0,
                max_characters
            )

            if split_at <= 0:
                split_at = max_characters

            pdf.drawString(
                left_margin,
                y,
                line[:split_at]
            )

            y -= line_height
            line = line[split_at:].lstrip()

            if y < bottom_margin:
                pdf.showPage()
                y = page_height - top_margin

        pdf.drawString(left_margin, y, line)
        y -= line_height

        if y < bottom_margin:
            pdf.showPage()
            y = page_height - top_margin

    pdf.save()


def convert_pdf_to_docx(file_path, output_file_path):
    converter = Converter(file_path)

    try:
        converter.convert(output_file_path)
    finally:
        converter.close()
